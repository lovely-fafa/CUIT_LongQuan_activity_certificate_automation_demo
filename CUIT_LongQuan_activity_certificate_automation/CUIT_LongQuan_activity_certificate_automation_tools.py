#!/usr/bin/python 3.10
# -*- coding: utf-8 -*- 
#
# @Time    : 2022-10-28 14:11
# @Author  : 发发
# @QQ      : 1315337973
# @File    : CUIT_LongQuan_activity_certificate_automation_tools.py
# @Software: PyCharm

import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from docxcompose.composer import Composer


def one_page_and_save_file(one_page_table_data: list[list], file_name: str, **kwargs):
    """
    one_page_table_data: 12行
    file_name:
    file_index:

    """
    document = Document('./static/template.docx')

    # 改第一段
    replace_str = document.paragraphs[1].text
    # 改字
    replace_str = re.sub(r'(X年X月X日X点至X点\(活动举行时间\))', kwargs['activity_date'], replace_str)
    replace_str = re.sub(r'(XX（组织/协会）)', kwargs['organization_name'], replace_str)
    replace_str = re.sub(r'(XXX（活动名称）)', kwargs['activity_name'], replace_str)
    document.paragraphs[1].text = replace_str
    # 奇奇怪怪的改字体的方法
    for one_run in document.paragraphs[1].runs:
        one_run.font.name = u'仿宋'
        one_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 中英文相同时，西文介个样子写？
        one_run.font.size = Pt(14)

    # 最后一段
    document.paragraphs[-1].text = kwargs['inscribe_date']
    for one_run in document.paragraphs[-1].runs:
        one_run.font.name = u'宋体'
        one_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        one_run.font.size = 152400

    # 取到第一个，也就是唯一一个表格
    table = document.tables[0]
    for row_index, row_data_list in enumerate(one_page_table_data):
        for col_index, col_text in enumerate(row_data_list):

            # if row_index == 0:
            #     continue
            cell_boj = table.cell(row_index + 1, col_index)

            # 介个格子的字体  (感觉介个地方写的不太好)
            cell_boj.text = one_page_table_data[row_index][col_index].strip()
            one_cell_runs = table.cell(row_index + 1, col_index).paragraphs[0].runs[0]  # 嘿嘿，第一次在学习一个库时用 dir 函数
            one_cell_runs.font.name = u'仿宋'
            one_cell_runs._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            one_cell_runs.font.size = 177800

            # 介个格子的段落
            # 居中
            cell_boj.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table.cell(row_index + 1, col_index).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 断后距
            table.cell(row_index + 1, col_index).paragraphs[0].paragraph_format.space_after = Pt(0)

    # 删空行
    while True:
        if len(table.rows) == len(one_page_table_data) + 1:
            break
        # print(del_row_index)
        row = table.rows[len(one_page_table_data) + 1]
        row._element.getparent().remove(row._element)

    document.save(file_name)


if __name__ == '__main__':
    test_data = [
        ['张三', '计算机学院', '222222222'],
        ['李四', '计院', '1'],
        ['王麻子', '统院', '1'],
    ]
    test_dict = {'activity_date': '2022年10月22日18:00至10月23日21:30',
                 'organization_name': '共青团成都信息工程大学委员会学生社团管理部',
                 'activity_name': '全国大学生数学建模竞赛',
                 'inscribe_date': '2022年2月31日'}
    one_page_and_save_file(test_data, '0', **test_dict)
